# agent_d_app.py
# Path: C:\Users\saisarkar\crewai_project
# Agent D — Streamlit + CrewAI-style orchestration, Gemini for LLM/RAG, Excel-all-tabs support
# UI: from Code 1 (header + center toggle + sidebars)

import io
import os
import re
import json
import glob
import uuid
import hashlib
import streamlit as st
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple

import streamlit as st
import pandas as pd
from dateutil.parser import parse as parse_date

# -- Greeting Utilities --
GREETINGS = ("hi", "hello", "hey", "good morning", "good afternoon", "good evening", "howdy", "greetings")

def is_user_greeting(text):
    norm = text.lower()
    return any(re.match(rf"^{g}\b", norm) for g in GREETINGS)

def friendly_greeting_reply(user_text):
    # Optionally extract name, but default to a friendly short reply
    return "Hello! 👋 How can I help you today?"

# ---- Optional libs
try:
    import faiss
except Exception:
    faiss = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from PIL import Image
except Exception:
    Image = None

# Optional local OCR (offline)
try:
    import pytesseract
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

try:
    from duckduckgo_search import DDGS
except Exception:
    DDGS = None

try:
    import trafilatura
except Exception:
    trafilatura = None

try:
    from docx import Document
except Exception:
    Document = None

# ---- Gemini (LLM)
import google.generativeai as genai

# ---- CrewAI (structure only; heavy lifting stays deterministic)
from crewai import Agent, Task, Crew

# -----------------------------
# CONSTANTS & PATHS
# -----------------------------
APP_TITLE = "Agent D"
DEFAULT_MODEL = "gemini-2.0-flash-lite-001"
EMBED_MODEL = "text-embedding-004"

BASE_DIR = os.getcwd()
DATA_DIR = os.path.join(BASE_DIR, "data_uploads")
REPO_DIR = os.path.join(BASE_DIR, "repository")
CHATS_DIR = os.path.join(BASE_DIR, "chats")
for d in (DATA_DIR, REPO_DIR, CHATS_DIR):
    os.makedirs(d, exist_ok=True)

MAX_TOTAL_UPLOAD_MB = 20
IMAGE_EXTS = (".png", ".jpg", ".jpeg")
EXCEL_EXTS = (".xlsx", ".xls", ".xlsm")

# -----------------------------
# PAGE & THEME (UI from Code 1)
# -----------------------------
st.set_page_config(page_title=APP_TITLE, page_icon="🤖", layout="wide")

st.markdown(
    """
<style>
h1, .stMarkdown h1 { font-weight: 700; letter-spacing: 0.3px; }
/* card support (kept from Code 1, mostly unused now) */
.card {
  border-radius: 14px; border: 1px solid rgba(120,120,120,0.16);
  padding: 14px 16px; background: rgba(250,250,250,0.65);
  box-shadow: 0 2px 6px rgba(0,0,0,0.03);
}
.toggle-row { display: flex; align-items: center; gap: 10px; }
.badge {
  display:inline-block; padding:2px 8px; border-radius:10px;
  font-size:12px; border:1px solid rgba(120,120,120,0.2); color:#444;
}
</style>
""",
    unsafe_allow_html=True,
)

# -----------------------------
# GEMINI INIT
# -----------------------------
genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
MODEL = st.secrets.get("MODEL", DEFAULT_MODEL)

# -----------------------------
# SESSION STATE
# -----------------------------
def ensure_state():
    if "internet_mode" not in st.session_state:
        st.session_state.internet_mode = False
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "file_store" not in st.session_state:
        # {checksum: rec}
        st.session_state.file_store = {}
    if "repo_index" not in st.session_state:
        # repo_index holds both tabular and text for RAG, plus images
        st.session_state.repo_index = {
            "tabular": [],
            "texts": [],
            "index": None,
            "metas": [],
            "images": [],
        }
    if "active_checksum" not in st.session_state:
        st.session_state.active_checksum = None
    if "chat_id" not in st.session_state:
        st.session_state.chat_id = None

ensure_state()

# -----------------------------
# GENERIC HELPERS
# -----------------------------
def file_checksum(b: bytes) -> str:
    h = hashlib.sha256()
    h.update(b)
    return h.hexdigest()

def safe_parse_date(text: str) -> Optional[str]:
    try:
        return parse_date(text).date().isoformat()
    except Exception:
        return None

def gemini_complete(system_prompt: str, user_prompt: str) -> str:
    model = genai.GenerativeModel(MODEL, system_instruction=system_prompt)
    resp = model.generate_content(user_prompt)
    return (resp.text or "").strip()

def gemini_vision_answer(prompt: str, image_bytes_list: List[bytes]) -> str:
    """Online multimodal answer from images (requires internet)."""
    if not image_bytes_list:
        return "No image provided."
    parts = [{"mime_type": "image/png", "data": b} for b in image_bytes_list]
    model = genai.GenerativeModel(MODEL)
    resp = model.generate_content([prompt] + parts)
    return (resp.text or "").strip()

def load_pdf_text(file_bytes: bytes) -> str:
    if PdfReader is None:
        return ""
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for pg in reader.pages:
        try:
            pages.append(pg.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n".join(pages)

def excel_to_frames(raw: bytes) -> List[pd.DataFrame]:
    """
    Robust Excel loader: supports all sheets, all Excel types.
    Returns list of dataframes with __sheet__ column.
    """
    xls = pd.ExcelFile(io.BytesIO(raw))
    frames: List[pd.DataFrame] = []
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        if not df.empty:
            df["__sheet__"] = sheet
            frames.append(df)
    if not frames:
        frames = [pd.read_excel(io.BytesIO(raw))]
    return frames

def df_to_text(df: pd.DataFrame, filename: str, sheet: str, max_rows: int = 40) -> str:
    """
    Convert a dataframe to a text snippet so that any kind of prompt
    can be answered via RAG. Limiting rows for token safety.
    """
    head = df.head(max_rows)
    # Use CSV-like representation, easier for model to parse
    return f"FILE: {filename}\nSHEET: {sheet}\n\n{head.to_csv(index=False)}"

def normalize_tabular(df: pd.DataFrame) -> pd.DataFrame:
    """
    Normalizes tabular data for allocation-style analytics,
    preserving all original columns. Works for ANY Excel sheet.
    """
    lower = {c.lower(): c for c in df.columns}

    def pick(opts):
        for o in opts:
            if o in lower:
                return lower[o]
        return None

    name_col = pick(["name", "provider", "employee", "resource"])
    alloc_col = pick(
        [
            "allocation_pct",
            "allocated %",
            "allocation",
            "allocation%",
            "utilization",
            "percent",
            "percentage",
        ]
    )
    date_col = pick(["date", "day", "as_of_date", "dt"])

    if name_col is None:
        name_col = df.columns[0]
    if alloc_col is None:
        numeric_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
        alloc_col = (
            numeric_cols[0]
            if numeric_cols
            else (df.columns[1] if len(df.columns) > 1 else df.columns[0])
        )

    out = pd.DataFrame()
    out["name"] = df[name_col].astype(str)
    out["allocation_pct"] = pd.to_numeric(df[alloc_col], errors="coerce")
    out["date"] = (
        pd.to_datetime(df[date_col], errors="coerce").dt.date if date_col else pd.NaT
    )

    for c in df.columns:
        if c not in [name_col, alloc_col, date_col]:
            out[c] = df[c]
    if "__sheet__" in df.columns:
        out["sheet"] = df["__sheet__"]
    return out

def chunk_text(text: str, max_chars: int = 1500) -> List[str]:
    paras = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    if not paras:
        return [text[:max_chars]]
    chunks: List[str] = []
    for p in paras:
        if len(p) <= max_chars:
            chunks.append(p)
        else:
            for i in range(0, len(p), max_chars):
                chunks.append(p[i : i + max_chars])
    return chunks

def ocr_image_bytes(img_bytes: bytes) -> str:
    if not OCR_AVAILABLE or Image is None:
        return ""
    try:
        img = Image.open(io.BytesIO(img_bytes))
        return pytesseract.image_to_string(img) or ""
    except Exception:
        return ""

def make_faiss_index(texts: List[str], metas: Optional[List[Dict[str, Any]]] = None):
    if faiss is None or not texts:
        return None, metas or []
    import numpy as np

    vecs = []
    for t in texts:
        emb = genai.embed_content(model=EMBED_MODEL, content=t)
        vecs.append(emb["embedding"])
    mat = np.array(vecs, dtype="float32")
    faiss.normalize_L2(mat)
    index = faiss.IndexFlatIP(mat.shape[1])
    index.add(mat)
    return index, metas or [{} for _ in texts]

def faiss_search(index, query: str, metas, k: int = 5):
    if faiss is None or index is None:
        return []
    import numpy as np

    q = genai.embed_content(model=EMBED_MODEL, content=query)["embedding"]
    q = np.array([q], dtype="float32")
    faiss.normalize_L2(q)
    D, I = index.search(q, k)
    out = []
    for i, score in zip(I[0], D[0]):
        if i == -1:
            continue
        md = metas[i] if metas and i < len(metas) else {}
        out.append((i, float(score), md))
    return out

# --- keyword-based retrieval (from Code 2) as fallback if FAISS unavailable
def keyword_retrieve(query: str, texts: List[str], metas: List[Dict[str, Any]], k: int = 4):
    if not texts:
        return []
    qw = set(query.lower().split())
    scored = []
    for i, t in enumerate(texts):
        score = len(qw & set(t.lower().split()))
        scored.append((score, i))
    scored.sort(key=lambda x: x[0], reverse=True)
    selected = [i for s, i in scored[:k] if s > 0]
    if not selected:
        selected = [0]  # fallback first chunk
    hits = []
    for i in selected:
        md = metas[i] if metas and i < len(metas) else {}
        hits.append((i, 1.0, md))
    return hits

# -----------------------------
# INDEXERS (Repository & Uploads)
# -----------------------------
def index_repository(repo_dir: str) -> Dict[str, Any]:
    """
    Build combined index for repository:
    - tabular_frames: for structured queries
    - text_chunks: from PDF, DOCX, TXT, Excel (all sheets), and OCRed images
    """
    tabular_frames: List[pd.DataFrame] = []
    text_chunks: List[str] = []
    metas: List[Dict[str, Any]] = []
    images_store: List[Dict[str, Any]] = []

    for path in glob.glob(os.path.join(repo_dir, "**/*"), recursive=True):
        if not os.path.isfile(path):
            continue
        low = path.lower()
        try:
            if low.endswith(EXCEL_EXTS):
                raw = open(path, "rb").read()
                excel_frames = excel_to_frames(raw)
                for df in excel_frames:
                    # tabular version
                    tabular_frames.append(normalize_tabular(df))
                    # textual version for RAG
                    sheet = df["__sheet__"].iloc[0] if "__sheet__" in df.columns else "Sheet"
                    text = df_to_text(df, os.path.basename(path), sheet)
                    for j, ch in enumerate(chunk_text(text)):
                        text_chunks.append(ch)
                        metas.append({"source": path, "chunk": j, "type": "excel"})
            elif low.endswith(".csv"):
                df = pd.read_csv(path)
                tabular_frames.append(normalize_tabular(df))
                text = df_to_text(df, os.path.basename(path), "csv")
                for j, ch in enumerate(chunk_text(text)):
                    text_chunks.append(ch)
                    metas.append({"source": path, "chunk": j, "type": "csv"})
            elif low.endswith(".pdf"):
                raw = open(path, "rb").read()
                txt = load_pdf_text(raw)
                for j, ch in enumerate(chunk_text(txt)):
                    text_chunks.append(ch)
                    metas.append({"source": path, "chunk": j, "type": "pdf"})
            elif low.endswith(".txt"):
                txt = open(path, "r", encoding="utf-8", errors="ignore").read()
                for j, ch in enumerate(chunk_text(txt)):
                    text_chunks.append(ch)
                    metas.append({"source": path, "chunk": j, "type": "txt"})
            elif low.endswith(".docx") and Document is not None:
                doc = Document(path)
                txt = "\n".join(p.text for p in doc.paragraphs)
                for j, ch in enumerate(chunk_text(txt)):
                    text_chunks.append(ch)
                    metas.append({"source": path, "chunk": j, "type": "docx"})
            elif low.endswith(IMAGE_EXTS):
                b = open(path, "rb").read()
                images_store.append({"path": path, "bytes": b})
                ocr_text = ocr_image_bytes(b)
                if ocr_text.strip():
                    for j, ch in enumerate(chunk_text(ocr_text)):
                        text_chunks.append(ch)
                        metas.append(
                            {"source": path, "chunk": j, "type": "image-ocr"}
                        )
        except Exception:
            continue

    combined_tabular = (
        pd.concat(tabular_frames, ignore_index=True) if tabular_frames else pd.DataFrame()
    )
    index, metas = make_faiss_index(text_chunks, metas) if text_chunks else (None, [])
    return {
        "tabular": [combined_tabular] if not combined_tabular.empty else [],
        "texts": text_chunks,
        "index": index,
        "metas": metas,
        "images": images_store,
    }

def add_uploaded_files(files) -> Tuple[Dict[str, Dict[str, Any]], List[str], int]:
    """
    Multi-file ingestion for uploads, mirroring repository behavior.
    """
    store_updates: Dict[str, Dict[str, Any]] = {}
    checksums: List[str] = []
    total_bytes = 0

    for up in files:
        total_bytes += up.size
        raw = up.read()
        checksum = file_checksum(raw)
        checksums.append(checksum)

        rec: Dict[str, Any] = {
            "filename": up.name,
            "uploaded_at": datetime.now().isoformat(timespec="seconds"),
        }
        kind: Optional[str] = None
        low = up.name.lower()

        try:
            if low.endswith(EXCEL_EXTS):
                frames = excel_to_frames(raw)
                tab_frames = []
                text_chunks = []
                metas = []
                for df in frames:
                    tab_frames.append(normalize_tabular(df))
                    sheet = df["__sheet__"].iloc[0] if "__sheet__" in df.columns else "Sheet"
                    text = df_to_text(df, up.name, sheet)
                    for j, ch in enumerate(chunk_text(text)):
                        text_chunks.append(ch)
                        metas.append(
                            {"source": up.name, "chunk": j, "type": "excel-upload"}
                        )
                if tab_frames:
                    rec["df"] = pd.concat(tab_frames, ignore_index=True)
                    kind = "tabular"
                if text_chunks:
                    idx, metas = make_faiss_index(text_chunks, metas)
                    rec.update({"texts": text_chunks, "index": idx, "metas": metas})
            elif low.endswith(".csv"):
                df = pd.read_csv(io.BytesIO(raw))
                rec["df"] = normalize_tabular(df)
                kind = "tabular"
                text = df_to_text(df, up.name, "csv-upload")
                t_chunks = chunk_text(text)
                idx, metas = make_faiss_index(
                    t_chunks,
                    [
                        {"source": up.name, "chunk": i, "type": "csv-upload"}
                        for i, _ in enumerate(t_chunks)
                    ],
                )
                rec.update({"texts": t_chunks, "index": idx, "metas": metas})
            elif low.endswith(".pdf"):
                txt = load_pdf_text(raw)
                chunks = chunk_text(txt)
                idx, metas = make_faiss_index(
                    chunks,
                    [{"source": up.name, "chunk": i, "type": "pdf-upload"} for i, _ in enumerate(chunks)],
                )
                rec.update({"texts": chunks, "index": idx, "metas": metas})
                kind = "text"
            elif low.endswith(".txt"):
                txt = raw.decode("utf-8", errors="ignore")
                chunks = chunk_text(txt)
                idx, metas = make_faiss_index(
                    chunks,
                    [{"source": up.name, "chunk": i, "type": "txt-upload"} for i, _ in enumerate(chunks)],
                )
                rec.update({"texts": chunks, "index": idx, "metas": metas})
                kind = "text"
            elif low.endswith(".docx") and Document is not None:
                doc = Document(io.BytesIO(raw))
                txt = "\n".join(p.text for p in doc.paragraphs)
                chunks = chunk_text(txt)
                idx, metas = make_faiss_index(
                    chunks,
                    [{"source": up.name, "chunk": i, "type": "docx-upload"} for i, _ in enumerate(chunks)],
                )
                rec.update({"texts": chunks, "index": idx, "metas": metas})
                kind = "text"
            elif low.endswith(IMAGE_EXTS):
                rec["image_bytes"] = raw
                kind = "image"
                ocr_text = ocr_image_bytes(raw)
                if ocr_text.strip():
                    chunks = chunk_text(ocr_text)
                    idx, metas = make_faiss_index(
                        chunks,
                        [{"source": up.name, "chunk": i, "type": "image-ocr-upload"} for i, _ in enumerate(chunks)],
                    )
                    rec.update({"texts": chunks, "index": idx, "metas": metas})
        except Exception as e:
            st.error(f"Ingestion error for {up.name}: {e}")

        if kind:
            rec["kind"] = kind
            store_updates[checksum] = rec

    return store_updates, checksums, total_bytes

# -----------------------------
# NL PARSER (strict JSON)
# -----------------------------
PARSER_SYSTEM = """
Translate a single user question into STRICT JSON:
Keys:
- intent: one of ["get_under_allocated","get_over_allocated","get_summary","tabular_select","rag_qa","web_qa","image_qa"]
- start_date: "YYYY-MM-DD" or null
- end_date: "YYYY-MM-DD" or null
- threshold: number or null
- columns: array or null
- conditions: array of {column, op, value} or null
Rules:
- Output ONLY the JSON object.
- Allocation/threshold language -> allocation intents.
- Generic table filters -> tabular_select.
- If the user mentions 'image' or 'this picture/photo/screenshot', use "image_qa".
- Current events/world facts -> web_qa (only if Internet mode ON).
"""

def nl_parse(user_text: str, internet_mode: bool) -> Dict[str, Any]:
    raw = gemini_complete(PARSER_SYSTEM, user_text)
    try:
        parsed = json.loads(raw)
    except Exception:
        m = re.search(r"\{.*\}", raw, re.S)
        parsed = (
            json.loads(m.group(0))
            if m
            else {
                "intent": "rag_qa",
                "start_date": None,
                "end_date": None,
                "threshold": None,
                "columns": None,
                "conditions": None,
            }
        )
    for k in ("start_date", "end_date"):
        v = parsed.get(k)
        if v:
            parsed[k] = safe_parse_date(v)
    if parsed.get("intent") == "web_qa" and not internet_mode:
        parsed["intent"] = "rag_qa"
    return parsed

# -----------------------------
# EXECUTORS (Tabular, RAG, Web, Image)
# -----------------------------
def run_tabular_allocation(
    df: pd.DataFrame,
    intent: str,
    start: Optional[str],
    end: Optional[str],
    threshold: Optional[float],
) -> Dict[str, Any]:
    work = df.copy()
    work["allocation_pct"] = pd.to_numeric(work["allocation_pct"], errors="coerce")
    if start:
        work = work[work["date"] >= pd.to_datetime(start).date()]
    if end:
        work = work[work["date"] <= pd.to_datetime(end).date()]

    if intent == "get_under_allocated":
        thr = threshold if threshold is not None else 100.0
        work = work[work["allocation_pct"] < thr]
        descriptor = f"allocation_pct < {thr}"
        work = work.sort_values("allocation_pct", ascending=True)
    else:
        thr = threshold if threshold is not None else 100.0
        work = work[work["allocation_pct"] > thr]
        descriptor = f"allocation_pct > {thr}"
        work = work.sort_values("allocation_pct", descending=False)
    q = f"date BETWEEN {start or '-inf'} AND {end or '+inf'} AND {descriptor}"
    return {"mode": "table", "table": work, "query": q, "count": int(work.shape[0])}

def apply_condition(df: pd.DataFrame, column: str, op: str, value: Any) -> pd.Series:
    op = (op or "=").strip().lower()
    if column not in df.columns:
        return pd.Series([False] * len(df))
    col = df[column]
    if op in ("=", "=="):
        return col.astype(str) == str(value)
    if op in ("!=", "<>"):
        return col.astype(str) != str(value)
    if op in (">", ">=", "<", "<="):
        left = pd.to_numeric(col, errors="coerce")
        try:
            val = float(value)
        except Exception:
            return pd.Series([False] * len(df))
        if op == ">":
            return left > val
        if op == ">=":
            return left >= val
        if op == "<":
            return left < val
        if op == "<=":
            return left <= val
    if op == "in" and isinstance(value, list):
        return col.astype(str).isin([str(x) for x in value])
    if op == "contains":
        return col.astype(str).str.contains(str(value), case=False, na=False)
    return pd.Series([False] * len(df))

def run_tabular_select(
    df: pd.DataFrame, conditions: Optional[List[Dict[str, Any]]]
) -> Dict[str, Any]:
    work = df.copy()
    mask = pd.Series([True] * len(work))
    if conditions:
        for cond in conditions:
            col, op, val = cond.get("column"), cond.get("op", "="), cond.get("value")
            mask = mask & apply_condition(work, col, op, val)
    work = work[mask]
    q_str = (
        " AND ".join(
            [f"{c.get('column')} {c.get('op')} {c.get('value')}" for c in (conditions or [])]
        )
        or "ALL"
    )
    return {"mode": "table", "table": work, "query": q_str, "count": int(work.shape[0])}

def local_rag(query: str, repo_idx: Dict[str, Any]) -> Dict[str, Any]:
    texts, metas, index = (
        repo_idx.get("texts"),
        repo_idx.get("metas"),
        repo_idx.get("index"),
    )
    if not texts:
        return {"mode": "error", "error": "No local text index configured."}

    # If FAISS is available, use it; otherwise fallback to keyword retrieval (from Code 2)
    if faiss is not None and index is not None:
        hits = faiss_search(index, query, metas, k=5)
    else:
        hits = keyword_retrieve(query, texts, metas, k=5)

    retrieved = []
    for i, score, md in hits:
        retrieved.append({"text": texts[i][:5000], "score": score, "meta": md})

    ctx = "\n\n".join([f"[Chunk {j+1}] {r['text']}" for j, r in enumerate(retrieved)])
    prompt = (
        "You are Agent D, a highly accurate analyst. Use ONLY the context below "
        "to answer if possible.\n\n"
        f"Context:\n{ctx}\n\nQuestion: {query}\nAnswer:"
    )
    answer = gemini_complete(
        "Be concise and grounded in the provided context. If unsure, say you cannot find that in the project data.",
        prompt,
    )
    return {"mode": "rag", "answer": answer, "sources": retrieved}

def fetch_web_context(query: str, n: int = 5) -> List[Dict[str, Any]]:
    pages = []
    if DDGS is None or trafilatura is None:
        return pages
    try:
        with DDGS() as ddgs:
            for r in ddgs.text(
                query, region="wt-wt", safesearch="moderate", max_results=n
            ):
                url = r.get("href") or r.get("url") or r.get("link")
                title = r.get("title") or r.get("body") or ""
                if not url:
                    continue
                html = trafilatura.fetch_url(url)
                text = trafilatura.extract(html) if html else ""
                if text:
                    pages.append({"url": url, "title": title, "text": text})
    except Exception:
        pass
    return pages

def web_qa(query: str) -> Dict[str, Any]:
    pages = fetch_web_context(query, 5)
    if not pages:
        return {"mode": "error", "error": "Web fetch returned no readable pages."}
    ctx = "\n\n---\n\n".join(
        [f"[{p['title']}] {p['text'][:2000]}" for p in pages[:5]]
    )
    answer = gemini_complete(
        "Answer ONLY with the provided web context; be precise and neutral.",
        f"{ctx}\n\nQuestion: {query}\nAnswer:",
    )
    return {
        "mode": "web",
        "answer": answer,
        "sources": [{"title": p["title"], "url": p["url"]} for p in pages],
    }

def image_qa_router(
    prompt: str, store: Dict[str, Dict[str, Any]], repo_idx: Dict[str, Any]
) -> Dict[str, Any]:
    """
    If internet ON -> Gemini multimodal; else try OCR text RAG.
    """
    image_bytes_list: List[bytes] = []
    for rec in store.values():
        if rec.get("kind") == "image" and rec.get("image_bytes"):
            image_bytes_list.append(rec["image_bytes"])
    for img in repo_idx.get("images", []):
        if img.get("bytes"):
            image_bytes_list.append(img["bytes"])

    if st.session_state.internet_mode:
        if not image_bytes_list:
            return {
                "mode": "image",
                "answer": "No images available. Upload or put in repository.",
            }
        ans = gemini_vision_answer(prompt, image_bytes_list[:2])
        return {
            "mode": "image",
            "answer": ans,
            "sources": [{"type": "image", "count": len(image_bytes_list)}],
        }
    else:
        # Offline: rely on OCR text already added to local index
        return local_rag(prompt, repo_idx)

# -----------------------------
# CREWAI STRUCTURE (Agents/Crew)
# -----------------------------
nl_agent = Agent(
    role="NLParserAgent",
    goal="Parse question to strict JSON.",
    backstory="Precise parameter extractor.",
    allow_delegation=False,
)
tabular_agent = Agent(
    role="TabularAgent",
    goal="Deterministic pandas answers for tables.",
    backstory="Analyst.",
    allow_delegation=False,
)
rag_agent = Agent(
    role="RAGAgent",
    goal="Local RAG with FAISS / keyword retrieval.",
    backstory="Grounded QA.",
    allow_delegation=False,
)
web_agent = Agent(
    role="WebAgent",
    goal="Web search and summarization.",
    backstory="Finds reliable info.",
    allow_delegation=False,
)
img_agent = Agent(
    role="ImageAgent",
    goal="Answer from images (multimodal or OCR).",
    backstory="Vision specialist.",
    allow_delegation=False,
)

crew = Crew(agents=[nl_agent, tabular_agent, rag_agent, web_agent, img_agent], tasks=[])

def crew_answer(user_text: str) -> Dict[str, Any]:
    """
    Main router: chooses between tabular analytics, web, rag, image.
    Uses combined repo + uploaded Excel frames for any "data" questions.
    """
    parsed = nl_parse(user_text, st.session_state.internet_mode)

    # Merge tabular sources (repo + active upload) for ANY Excel/CSV queries
    tabular_sources: List[pd.DataFrame] = []
    if st.session_state.repo_index.get("tabular"):
        tabular_sources.extend(st.session_state.repo_index["tabular"])
    if (
        st.session_state.active_checksum
        and st.session_state.active_checksum in st.session_state.file_store
    ):
        rec = st.session_state.file_store[st.session_state.active_checksum]
        if rec.get("kind") == "tabular" and isinstance(rec.get("df"), pd.DataFrame):
            tabular_sources.append(rec["df"])

    combined_df = (
        pd.concat(tabular_sources, ignore_index=True) if tabular_sources else pd.DataFrame()
    )

    intent = parsed.get("intent")

    if intent in {"get_under_allocated", "get_over_allocated"} and not combined_df.empty:
        res = run_tabular_allocation(
            combined_df,
            intent,
            parsed.get("start_date"),
            parsed.get("end_date"),
            parsed.get("threshold"),
        )
        return {"mode": "tabular", "parsed": parsed, "result": res}

    if intent == "get_summary" and not combined_df.empty:
        desc = combined_df["allocation_pct"].describe().to_frame().reset_index()
        return {
            "mode": "tabular",
            "parsed": parsed,
            "result": {
                "mode": "summary",
                "table": desc,
                "query": "SUMMARY(allocation_pct)",
                "count": int(combined_df.shape[0]),
            },
        }

    if intent == "tabular_select" and not combined_df.empty:
        res = run_tabular_select(combined_df, parsed.get("conditions"))
        return {"mode": "tabular", "parsed": parsed, "result": res}

    if intent == "web_qa" and st.session_state.internet_mode:
        res = web_qa(user_text)
        return {"mode": "web", "parsed": parsed, "result": res}

    if intent == "image_qa":
        res = image_qa_router(
            user_text, st.session_state.file_store, st.session_state.repo_index
        )
        return {"mode": "image", "parsed": parsed, "result": res}

    # default → local RAG (Excel sheets, PDFs, DOCX, TXT, OCR, etc.)
    res = local_rag(user_text, st.session_state.repo_index)
    return {"mode": "rag", "parsed": parsed, "result": res}

# -----------------------------
# LAYOUT (UI from Code 1)
# -----------------------------
left, mid, right = st.columns([1, 2, 1])
with mid:
    st.markdown(
        """
        <div style="display:flex; justify-content:center; align-items:flex-end;">
            <span style="font-size:48px; font-weight:700; letter-spacing:0.3px; color:#2e2f38; line-height:1;">
                Agent D
                <span style="
                    display:inline-block;
                    width:18px;
                    height:18px;
                    background: #00FF00;
                    border-radius:50%;
                    margin-left:2px;
                    position:relative;
                    top:-2px;
                " title="Agent is active"></span>
            </span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Center the toggle below the header
    st.markdown(
        '<div style="display:flex;justify-content:center;margin-top:20px;margin-bottom:10px;">',
        unsafe_allow_html=True,
    )
    st.session_state.internet_mode = st.toggle(
        "🌐",
        value=st.session_state.internet_mode,
        help="ON = may use internet for answers. OFF = strictly local files/repo.",
    )
    st.markdown("</div>", unsafe_allow_html=True)

# -----------------------------
# SIDEBAR — LEFT: BROWSE FILES
# -----------------------------
with st.sidebar:
    st.subheader("📁 Browse files (upload)")
    uploads = st.file_uploader(
        "Drop files (XLSX/XLS/XLSM/CSV/PDF/TXT/DOCX/PNG/JPG)",
        type=["xlsx", "xls", "xlsm", "csv", "pdf", "txt", "docx", "png", "jpg", "jpeg"],
        accept_multiple_files=True,
        key="uploader_multi",
    )
    if uploads:
        total = sum([u.size for u in uploads])
        if total > MAX_TOTAL_UPLOAD_MB * 1024 * 1024:
            st.error(
                f"Total upload too large: {(total/1024/1024):.2f} MB (limit {MAX_TOTAL_UPLOAD_MB} MB)"
            )
        else:
            updates, checksums, _ = add_uploaded_files(uploads)
            st.session_state.file_store.update(updates)
            if checksums:
                st.session_state.active_checksum = checksums[-1]
            st.success(f"Added {len(checksums)} file(s).")

    if st.session_state.file_store:
        st.caption("Active upload")
        keys = list(st.session_state.file_store.keys())
        labels = [
            f"{st.session_state.file_store[k]['filename']} · {k[:8]}" for k in keys
        ]
        current = (
            keys.index(st.session_state.active_checksum)
            if st.session_state.active_checksum in keys
            else 0
        )
        choice = st.selectbox(
            "Select active:", options=list(range(len(keys))), format_func=lambda i: labels[i], index=current
        )
        st.session_state.active_checksum = keys[choice]

    st.markdown("---")
    st.subheader("📦 Repository")
    st.caption(f"Folder: `{REPO_DIR}`")
    if st.button("Re-index repository"):
        st.session_state.repo_index = index_repository(REPO_DIR)
        st.success("Repository re-indexed.")
    if (
        not st.session_state.repo_index.get("tabular")
        and not st.session_state.repo_index.get("texts")
        and not st.session_state.repo_index.get("images")
    ):
        st.info(
            "Add your PDFs/TXTs/CSVs/XLSXs/XLS/XLSM/DOCX/Images to the repository and click Re-index."
        )

# -----------------------------
# SIDEBAR — RIGHT: CHAT HISTORY
# -----------------------------
with st.sidebar:
    st.markdown("---")
    st.subheader("🗂️ Chat History")
    files = sorted(glob.glob(os.path.join(CHATS_DIR, "*.json")))
    names = ["(new)"] + [os.path.basename(p)[:-5] for p in files]
    opened = st.selectbox("Open:", options=names, index=0)
    if opened != "(new)":
        with open(os.path.join(CHATS_DIR, f"{opened}.json"), "r", encoding="utf-8") as f:
            data = json.load(f)
        st.session_state.chat_id = data.get("chat_id", opened)
        st.session_state.messages = data.get("messages", [])
        st.success(f"Loaded chat: {opened}")
    cols = st.columns(2)
    with cols[0]:
        if st.button("🧹 New chat"):
            st.session_state.chat_id = (
                datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:6]
            )
            st.session_state.messages = []
            st.success("Started a new chat.")
    with cols[1]:
        if st.session_state.messages and st.button("💾 Save chat"):
            if not st.session_state.chat_id:
                st.session_state.chat_id = (
                    datetime.now().strftime("%Y%m%d_%H%M%S")
                    + "_"
                    + uuid.uuid4().hex[:6]
                )
            out = {
                "chat_id": st.session_state.chat_id,
                "saved_at": datetime.now().isoformat(timespec="seconds"),
                "messages": st.session_state.messages,
            }
            with open(
                os.path.join(CHATS_DIR, f"{st.session_state.chat_id}.json"),
                "w",
                encoding="utf-8",
            ) as f:
                json.dump(out, f, ensure_ascii=False, indent=2)
            st.success(f"Saved chat: {st.session_state.chat_id}")

# -----------------------------
# MAIN CHAT (ChatGPT-style)
# -----------------------------
st.markdown("### 💬 Chat")

# show history
for m in st.session_state.messages:
    with st.chat_message(m["role"]):
        st.markdown(m["content"])

user_text = st.chat_input(
    "Ask Agent D anything about your project data or the web (depending on toggle)…"
)

if user_text:
    st.session_state.messages.append({"role": "user", "content": user_text})
    with st.chat_message("user"):
        st.markdown(user_text)

    result = crew_answer(user_text)

    with st.chat_message("assistant"):
        mode = result.get("mode")
        if mode == "tabular":
            res = result["result"]
            if res["mode"] == "summary":
                st.markdown("**Summary (allocation_pct)**")
                st.dataframe(res["table"])
                st.info(f"Rows considered: {res['count']}")
                st.session_state.messages.append(
                    {"role": "assistant", "content": "Posted summary stats."}
                )
            elif res["mode"] == "table":
                st.markdown("**Tabular results**")
                with st.expander("Query descriptor"):
                    st.code(res["query"], language="sql")
                st.dataframe(res["table"])
                csv = res["table"].to_csv(index=False).encode("utf-8")
                st.download_button(
                    "Download CSV",
                    data=csv,
                    file_name="agentd_result.csv",
                    mime="text/csv",
                )
                st.success(f"{res['count']} row(s).")
                st.session_state.messages.append(
                    {
                        "role": "assistant",
                        "content": f"Returned {res['count']} rows.",
                    }
                )
            else:
                st.warning("No results.")
                st.session_state.messages.append(
                    {"role": "assistant", "content": "No results."}
                )

        elif mode == "rag":
            res = result["result"]
            if res.get("mode") == "rag":
                st.markdown("**Local RAG Answer**")
                st.markdown(res["answer"])
                if res.get("sources"):
                    with st.expander("Sources"):
                        for j, s in enumerate(res["sources"], 1):
                            st.markdown(
                                f"**Chunk {j}** · score={s['score']:.3f} · source={s['meta'].get('source','')}"
                            )
                            text = s["text"]
                            st.write(
                                (text[:1000] + "…") if len(text) > 1000 else text
                            )
                st.session_state.messages.append(
                    {
                        "role": "assistant",
                        "content": res["answer"][:2000],
                    }
                )
            else:
                st.error(res.get("error", "Local RAG error"))
                st.session_state.messages.append(
                    {"role": "assistant", "content": "RAG error."}
                )

        elif mode == "web":
            res = result["result"]
            if res.get("mode") == "web":
                st.markdown("**Web-grounded Answer**")
                st.markdown(res["answer"])
                if res.get("sources"):
                    with st.expander("Sources"):
                        for s in res["sources"]:
                            st.markdown(f"- [{s['title']}]({s['url']})")
                st.session_state.messages.append(
                    {
                        "role": "assistant",
                        "content": res["answer"][:2000],
                    }
                )
            else:
                st.error(res.get("error", "Web QA error"))
                st.session_state.messages.append(
                    {"role": "assistant", "content": "Web QA error."}
                )

        elif mode == "image":
            res = result["result"]
            st.markdown("**Image Answer**")
            st.markdown(res.get("answer", "(no answer)"))
            st.session_state.messages.append(
                {
                    "role": "assistant",
                    "content": res.get("answer", "(no answer)")[:2000],
                }
            )

        else:
            st.error("Unknown mode.")
            st.session_state.messages.append(
                {"role": "assistant", "content": "Unknown mode."}
            )

    # autosave each turn
    if not st.session_state.chat_id:
        st.session_state.chat_id = (
            datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:6]
        )
    out = {
        "chat_id": st.session_state.chat_id,
        "saved_at": datetime.now().isoformat(timespec="seconds"),
        "messages": st.session_state.messages,
    }
    with open(
        os.path.join(CHATS_DIR, f"{st.session_state.chat_id}.json"),
        "w",
        encoding="utf-8",
    ) as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
